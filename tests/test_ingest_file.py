"""Integration tests for ingest_raw_text with real document files."""

from pathlib import Path
import sqlite3
import sys

ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "src"
if str(SRC) not in sys.path:
    sys.path.insert(0, str(SRC))

from brain4me.converters import convert_file
from brain4me.ingest import ingest_raw_text
from brain4me.storage import initialize_database
from tests.conftest import WorkspaceTempDirTestCase


class IngestFileTests(WorkspaceTempDirTestCase):
    """Integration tests: convert a file, ingest it, verify entities in SQLite."""

    def _create_txt_fixture(self) -> Path:
        path = self.temp_dir / "sample.txt"
        text = (
            "Projeto: Brain4me\n"
            "Problema: testar ingestao de TXT\n"
            "Decisao: usar SQLite no MVP\n"
            "Evidencia: reduz complexidade operacional\n"
            "Alternativa: Neo4j\n"
            "Risco: revisar escolha no futuro\n"
        )
        path.write_text(text, encoding="utf-8")
        return path

    def _create_docx_fixture(self) -> Path:
        from docx import Document
        path = self.temp_dir / "sample.docx"
        doc = Document()
        doc.add_paragraph("Projeto: Brain4me")
        doc.add_paragraph("Problema: testar ingestao de DOCX")
        doc.add_paragraph("Decisao: usar SQLite no MVP")
        doc.add_paragraph("Evidencia: setup mais simples")
        doc.add_paragraph("Alternativa: Neo4j")
        doc.add_paragraph("Risco: consultas relacionais avancadas")
        doc.save(str(path))
        return path

    def _create_pdf_fixture(self) -> Path:
        import fitz
        path = self.temp_dir / "sample.pdf"
        doc = fitz.open()
        page = doc.new_page()
        text = (
            "Projeto: Brain4me\n"
            "Problema: testar ingestao de PDF\n"
            "Decisao: usar SQLite no MVP\n"
            "Evidencia: reduz complexidade operacional\n"
            "Alternativa: Neo4j\n"
            "Risco: revisar escolha no futuro\n"
        )
        page.insert_text((50, 100), text, fontsize=10)
        doc.save(str(path))
        doc.close()
        return path

    # ------------------------------------------------------------------
    # TXT integration
    # ------------------------------------------------------------------
    def test_ingest_txt_creates_entities(self):
        db_path = self.temp_dir / "brain4me.db"
        initialize_database(db_path)
        txt_path = self._create_txt_fixture()

        result = ingest_raw_text(
            db_path=db_path,
            raw_text=txt_path.read_text(encoding="utf-8"),
            source_path=str(txt_path),
            source_type="txt",
            compartment="test",
        )

        self.assertGreaterEqual(result.entities_created, 4)
        self.assertGreaterEqual(result.relations_created, 2)
        self.assertGreaterEqual(result.context_nodes_created, 3)
        self.assertGreaterEqual(result.context_edges_created, 2)
        self.assertGreaterEqual(result.memory_entries_created, 1)

        with sqlite3.connect(db_path) as conn:
            types = {
                row[0]
                for row in conn.execute("SELECT entity_type FROM entities").fetchall()
            }
            self.assertIn("Project", types)
            self.assertIn("Decision", types)
            self.assertIn("Evidence", types)
            self.assertIn("Alternative", types)
            self.assertIn("Risk", types)

    # ------------------------------------------------------------------
    # DOCX integration
    # ------------------------------------------------------------------
    def test_ingest_docx_creates_entities(self):
        db_path = self.temp_dir / "brain4me.db"
        initialize_database(db_path)
        docx_path = self._create_docx_fixture()

        text = convert_file(docx_path)
        result = ingest_raw_text(
            db_path=db_path,
            raw_text=text,
            source_path=str(docx_path),
            source_type="docx",
            compartment="test",
        )

        self.assertGreaterEqual(result.entities_created, 4)
        self.assertGreaterEqual(result.relations_created, 2)

        with sqlite3.connect(db_path) as conn:
            types = {
                row[0]
                for row in conn.execute("SELECT entity_type FROM entities").fetchall()
            }
            self.assertIn("Project", types)
            self.assertIn("Decision", types)

    # ------------------------------------------------------------------
    # PDF integration
    # ------------------------------------------------------------------
    def test_ingest_pdf_creates_entities(self):
        db_path = self.temp_dir / "brain4me.db"
        initialize_database(db_path)
        pdf_path = self._create_pdf_fixture()

        text = convert_file(pdf_path)
        result = ingest_raw_text(
            db_path=db_path,
            raw_text=text,
            source_path=str(pdf_path),
            source_type="pdf",
            compartment="test",
        )

        self.assertGreaterEqual(result.entities_created, 4)
        self.assertGreaterEqual(result.relations_created, 2)

        with sqlite3.connect(db_path) as conn:
            types = {
                row[0]
                for row in conn.execute("SELECT entity_type FROM entities").fetchall()
            }
            self.assertIn("Project", types)
            self.assertIn("Decision", types)

    # ------------------------------------------------------------------
    # Edge cases
    # ------------------------------------------------------------------
    def test_ingest_raw_text_custom_title(self):
        db_path = self.temp_dir / "brain4me.db"
        initialize_database(db_path)
        txt_path = self._create_txt_fixture()

        result = ingest_raw_text(
            db_path=db_path,
            raw_text=txt_path.read_text(encoding="utf-8"),
            source_path=str(txt_path),
            source_type="txt",
            compartment="test",
            title="Meu titulo customizado",
        )

        self.assertGreaterEqual(result.entities_created, 1)

        with sqlite3.connect(db_path) as conn:
            notes = conn.execute(
                "SELECT summary FROM notes WHERE summary = ?",
                ("Meu titulo customizado",),
            ).fetchall()
            self.assertEqual(len(notes), 1)

    def test_ingest_empty_text(self):
        db_path = self.temp_dir / "brain4me.db"
        initialize_database(db_path)

        result = ingest_raw_text(
            db_path=db_path,
            raw_text="",
            source_path="empty.txt",
            source_type="txt",
            compartment="default",
        )

        # No entities extracted from empty text
        self.assertEqual(result.entities_created, 0)
        self.assertEqual(result.relations_created, 0)
        self.assertEqual(result.memory_entries_created, 0)

        # But a source and note should still be created
        with sqlite3.connect(db_path) as conn:
            source_count = conn.execute(
                "SELECT COUNT(*) FROM sources"
            ).fetchone()[0]
            self.assertEqual(source_count, 1)

"""Unit tests for the document converter functions."""

from pathlib import Path
import sys

ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "src"
if str(SRC) not in sys.path:
    sys.path.insert(0, str(SRC))

import unittest

from brain4me.converters import (
    convert_file,
    extract_text_from_docx,
    extract_text_from_pdf,
    extract_text_from_txt,
)


class ConverterTests(unittest.TestCase):
    """Test each extractor with real files created inline."""

    def setUp(self):
        self.temp_dir = ROOT / ".tmp-tests" / self.id().replace(".", "_")
        self.temp_dir.mkdir(parents=True, exist_ok=True)

    def tearDown(self):
        import shutil
        shutil.rmtree(self.temp_dir, ignore_errors=True)

    # ------------------------------------------------------------------
    # TXT
    # ------------------------------------------------------------------
    def test_extract_text_from_txt_utf8(self):
        path = self.temp_dir / "sample.txt"
        expected = "Projeto: Brain4me\nProblema: testar conversao\n"
        path.write_text(expected, encoding="utf-8")
        result = extract_text_from_txt(path)
        self.assertEqual(result, expected)

    def test_extract_text_from_txt_latin1_fallback(self):
        """Verify Latin-1 fallback works for non-UTF-8 content."""
        path = self.temp_dir / "latin1.txt"
        # \xe9 is 'é' in Latin-1 but invalid multi-byte UTF-8
        raw = b"Projeto: Brain4me\nDecis\xe3o: usar SQLite\n"
        path.write_bytes(raw)
        result = extract_text_from_txt(path)
        self.assertIn("Brain4me", result)
        self.assertIn("Decis", result)

    # ------------------------------------------------------------------
    # DOCX
    # ------------------------------------------------------------------
    def test_extract_text_from_docx(self):
        path = self.temp_dir / "sample.docx"
        self._create_minimal_docx(path)
        result = extract_text_from_docx(path)
        self.assertIn("Projeto: Brain4me", result)
        self.assertIn("Decisao: usar SQLite", result)
        self.assertIn("Evidencia: reduz complexidade", result)

    # ------------------------------------------------------------------
    # PDF
    # ------------------------------------------------------------------
    def test_extract_text_from_pdf(self):
        path = self.temp_dir / "sample.pdf"
        self._create_minimal_pdf(path)
        result = extract_text_from_pdf(path)
        self.assertIn("Projeto: Brain4me", result)
        self.assertIn("Decisao: usar SQLite", result)

    # ------------------------------------------------------------------
    # convert_file router
    # ------------------------------------------------------------------
    def test_convert_file_txt(self):
        path = self.temp_dir / "test.txt"
        path.write_text("Projeto: Brain4me\n", encoding="utf-8")
        text = convert_file(path)
        self.assertEqual(text, "Projeto: Brain4me\n")

    def test_convert_file_docx(self):
        path = self.temp_dir / "test.docx"
        self._create_minimal_docx(path)
        text = convert_file(path)
        self.assertIn("Projeto: Brain4me", text)

    def test_convert_file_pdf(self):
        path = self.temp_dir / "test.pdf"
        self._create_minimal_pdf(path)
        text = convert_file(path)
        self.assertIn("Projeto: Brain4me", text)

    def test_convert_file_unsupported_raises(self):
        path = self.temp_dir / "test.xyz"
        path.write_text("whatever", encoding="utf-8")
        with self.assertRaises(ValueError) as ctx:
            convert_file(path)
        self.assertIn("Unsupported file format", str(ctx.exception))

    # ------------------------------------------------------------------
    # helpers
    # ------------------------------------------------------------------
    def _create_minimal_docx(self, path: Path) -> None:
        from docx import Document
        doc = Document()
        doc.add_paragraph("Projeto: Brain4me")
        doc.add_paragraph("Problema: escolher persistencia")
        doc.add_paragraph("Decisao: usar SQLite")
        doc.add_paragraph("Evidencia: reduz complexidade")
        doc.add_paragraph("Risco: revisar no futuro")
        doc.save(str(path))

    def _create_minimal_pdf(self, path: Path) -> None:
        import fitz  # pymupdf
        doc = fitz.open()
        page = doc.new_page()
        text = (
            "Projeto: Brain4me\n"
            "Problema: escolher persistencia\n"
            "Decisao: usar SQLite\n"
            "Evidencia: reduz complexidade operacional\n"
            "Alternativa: Neo4j\n"
            "Risco: limitar consultas de grafo\n"
        )
        page.insert_text((50, 100), text, fontsize=10)
        doc.save(str(path))
        doc.close()
